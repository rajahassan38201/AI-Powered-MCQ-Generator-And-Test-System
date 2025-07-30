import streamlit as st
import os
import json
from io import BytesIO
from dotenv import load_dotenv
load_dotenv()  # Load environment variables from .env file

# Langchain imports
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.docstore.document import Document

# For DOCX handling (optional, requires python-docx)
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
    st.warning("`python-docx` not found. DOCX file uploads will not be supported. Install with `pip install python-docx`.")

# --- Configuration and Setup ---

# Set Streamlit page configuration
st.set_page_config(
    page_title="AI MCQ Generator",
    layout="wide",
    initial_sidebar_state="expanded"
)
#Add your Gemini API key in .env file

llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.7)

# --- Helper Functions ---

def get_file_text(uploaded_file):
    """
    Extracts text content from uploaded files (TXT, PDF, DOCX).
    """
    file_extension = uploaded_file.name.split('.')[-1].lower()
    text_content = ""

    with BytesIO(uploaded_file.getvalue()) as text_buffer:
        if file_extension == "txt":
            text_content = text_buffer.read().decode("utf-8")
        elif file_extension == "pdf":
            # Save to a temporary file for PyPDFLoader
            with open("temp.pdf", "wb") as f:
                f.write(text_buffer.read())
            loader = PyPDFLoader("temp.pdf")
            pages = loader.load()
            text_content = "\n".join([page.page_content for page in pages])
            os.remove("temp.pdf") # Clean up temporary file
        elif file_extension == "docx" and DocxDocument:
            doc = DocxDocument(text_buffer)
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        else:
            st.error(f"Unsupported file type: .{file_extension}. Please upload a TXT, PDF, or DOCX file.")
            return None
    return text_content

def generate_mcqs(text_content, num_mcqs):
    """
    Generates MCQs using Langchain and Google Gemini.
    Returns:
        tuple: (list of MCQs for display, dict of correct answers)
    """
    if not text_content:
        st.error("No text content available to generate MCQs from.")
        return [], {}

    # Chunk the text
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,  # Adjust chunk size based on model context window
        chunk_overlap=200,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
    )
    # Langchain expects a list of Document objects
    docs = [Document(page_content=text_content, metadata={"source": "uploaded_document"})]
    chunks = text_splitter.split_documents(docs)

    if not chunks:
        st.error("Could not split the document into readable chunks. Please ensure the document has sufficient content.")
        return [], {}

    # Define the prompt template for MCQ generation
    # We instruct the model to output JSON for easy parsing
    prompt_template = PromptTemplate(
        input_variables=["text_chunk", "num_mcqs"],
        template="""
        You are an expert educator tasked with creating multiple-choice questions (MCQs) from provided text.
        Generate exactly {num_mcqs} multiple-choice questions based on the following text chunk.
        Each question should have 4 options (A, B, C, D) and clearly indicate the correct answer.
        The options should be plausible distractors, and only one option should be correct.
        Ensure the questions directly relate to the content of the text.

        Output the MCQs in a JSON array format. Each MCQ object in the array should have the following keys:
        - "question": The question text.
        - "options": An array of 4 strings, each representing an option (e.g., ["A. Option 1", "B. Option 2", "C. Option 3", "D. Option 4"]).
        - "correct_answer": The full string of the correct option (e.g., "B. Paris").

        Example JSON format:
        [
            {{
                "question": "What is the capital of France?",
                "options": ["A. Berlin", "B. Paris", "C. Rome", "D. Madrid"],
                "correct_answer": "B. Paris"
            }},
            {{
                "question": "Which planet is known as the Red Planet?",
                "options": ["A. Earth", "B. Mars", "C. Jupiter", "D. Venus"],
                "correct_answer": "B. Mars"
            }}
        ]

        Text to generate MCQs from:
        ---
        {text_chunk}
        ---
        """
    )

    mcq_chain = LLMChain(llm=llm, prompt=prompt_template)
    
    combined_text_for_mcqs = chunks[0].page_content if chunks else ""
    
    display_mcqs = []
    correct_answers_map = {}

    if not combined_text_for_mcqs:
        return [], {}

    try:
        with st.spinner(f"Generating {num_mcqs} MCQs... This may take a moment."):
            response = mcq_chain.invoke({"text_chunk": combined_text_for_mcqs, "num_mcqs": num_mcqs})
            mcqs_json_str = response.get('text', '').strip()

            # Attempt to clean up common LLM formatting issues (e.g., ```json)
            if mcqs_json_str.startswith("```json"):
                mcqs_json_str = mcqs_json_str[len("```json"):].strip()
            if mcqs_json_str.endswith("```"):
                mcqs_json_str = mcqs_json_str[:-len("```")].strip()

            generated_mcqs = json.loads(mcqs_json_str)
            
            if not isinstance(generated_mcqs, list):
                raise ValueError("Expected a JSON array of MCQs.")
            
            for mcq in generated_mcqs:
                if not all(k in mcq for k in ["question", "options", "correct_answer"]):
                    raise ValueError("Each MCQ must have 'question', 'options', and 'correct_answer' keys.")
                if not isinstance(mcq["options"], list) or len(mcq["options"]) != 4:
                    raise ValueError("Options must be a list of 4 strings.")
                
                # Store correct answer separately
                correct_answers_map[mcq["question"]] = mcq["correct_answer"]
                
                # Create MCQ for display (without correct_answer)
                display_mcqs.append({
                    "question": mcq["question"],
                    "options": mcq["options"]
                })

    except json.JSONDecodeError as e:
        st.error(f"Failed to parse MCQs from AI response. Please try again or refine your document. Error: {e}")
        st.text(f"Raw AI response (for debugging): {mcqs_json_str}")
        display_mcqs = []
        correct_answers_map = {}
    except ValueError as e:
        st.error(f"Generated MCQs have an unexpected structure. Error: {e}")
        st.text(f"Raw AI response (for debugging): {mcqs_json_str}")
        display_mcqs = []
        correct_answers_map = {}
    except Exception as e:
        st.error(f"An error occurred during MCQ generation: {e}")
        display_mcqs = []
        correct_answers_map = {}

    return display_mcqs, correct_answers_map

# --- Streamlit UI ---

st.title("📚 AI-Powered MCQ Generator And Test System")
st.markdown("""
Welcome, Upload your document (TXT, PDF, or DOCX),
specify how many multiple-choice questions you need, and let AI generate them for you.
""")

# Initialize session state variables for test flow
if 'mcqs_for_display' not in st.session_state:
    st.session_state.mcqs_for_display = []
if 'correct_answers_map' not in st.session_state:
    st.session_state.correct_answers_map = {}
if 'user_answers' not in st.session_state:
    st.session_state.user_answers = {}
if 'test_submitted' not in st.session_state:
    st.session_state.test_submitted = False
if 'text_content' not in st.session_state:
    st.session_state.text_content = ""
if 'current_state' not in st.session_state:
    st.session_state.current_state = "upload" # States: "upload", "test", "results"

# Sidebar for inputs
with st.sidebar:
    st.header("Upload Document & Settings")
    uploaded_file = st.file_uploader(
        "Upload a file (TXT, PDF, DOCX)",
        type=["txt", "pdf", "docx"],
        help="Please upload a text-based document or book."
    )

    num_mcqs = st.number_input(
        "Number of MCQs to generate:",
        min_value=1,
        max_value=20,
        value=5,
        step=1,
        help="Enter the desired number of multiple-choice questions (1-20)."
    )

    generate_button = st.button("Generate MCQs", type="primary")
    
    # Reset button to go back to upload state
    if st.button("Start New Test"):
        st.session_state.mcqs_for_display = []
        st.session_state.correct_answers_map = {}
        st.session_state.user_answers = {}
        st.session_state.test_submitted = False
        st.session_state.text_content = ""
        st.session_state.current_state = "upload"
        st.rerun() # Rerun to clear the main content area

# --- Main Content Area Logic ---

if generate_button:
    if uploaded_file is None:
        st.warning("Please upload a document to generate MCQs.")
    elif num_mcqs <= 0:
        st.warning("Please enter a valid number of MCQs (greater than 0).")
    else:
        # Reset state for new generation
        st.session_state.mcqs_for_display = []
        st.session_state.correct_answers_map = {}
        st.session_state.user_answers = {}
        st.session_state.test_submitted = False
        st.session_state.text_content = ""
        st.session_state.current_state = "upload" # Temporarily set to upload during processing

        with st.spinner("Extracting text from document..."):
            text_content = get_file_text(uploaded_file)

        if text_content:
            st.session_state.text_content = text_content
            st.info(f"Text extracted successfully! Document size: {len(text_content)} characters.")

            mcqs_display, mcqs_correct_map = generate_mcqs(text_content, num_mcqs)
            
            if mcqs_display:
                st.session_state.mcqs_for_display = mcqs_display
                st.session_state.correct_answers_map = mcqs_correct_map
                st.session_state.user_answers = {mcq['question']: None for mcq in mcqs_display} # Initialize user answers
                st.session_state.current_state = "test"
                st.success(f"Successfully generated {len(mcqs_display)} MCQs! Please take the test below.")
            else:
                st.error("Failed to generate MCQs. Please check the document content and try again.")
                st.session_state.current_state = "upload"
        else:
            st.error("Could not extract text from the uploaded file. Please ensure it's a valid TXT, PDF, or DOCX.")
            st.session_state.current_state = "upload"

# Display test or results based on state
if st.session_state.current_state == "test" and st.session_state.mcqs_for_display:
    st.header("Take the MCQ Test")
    with st.form("mcq_test_form"):
        for i, mcq in enumerate(st.session_state.mcqs_for_display):
            st.markdown(f"**Question {i+1}:** {mcq['question']}")
            # Use a unique key for each radio button group
            user_selection = st.radio(
                f"Select an option for Q{i+1}:",
                mcq['options'],
                key=f"q_{i}_radio",
                index=None # No option selected by default
            )
            # Store user's answer in session state
            st.session_state.user_answers[mcq['question']] = user_selection
            st.markdown("---")
        
        submit_test_button = st.form_submit_button("Submit Test")

        if submit_test_button:
            # Check if all questions are answered
            unanswered_questions = [q for q, a in st.session_state.user_answers.items() if a is None]
            if unanswered_questions:
                st.warning("Please answer all questions before submitting.")
                # Keep state as "test" so user can complete
            else:
                st.session_state.test_submitted = True
                st.session_state.current_state = "results"
                st.rerun() # Rerun to switch to results view

elif st.session_state.current_state == "results" and st.session_state.test_submitted:
    st.header("Test Results")
    
    correct_count = 0
    total_questions = len(st.session_state.mcqs_for_display)
    results_data = []

    for mcq in st.session_state.mcqs_for_display:
        question = mcq['question']
        user_answer = st.session_state.user_answers.get(question)
        correct_answer = st.session_state.correct_answers_map.get(question)

        is_correct = (user_answer == correct_answer)
        if is_correct:
            correct_count += 1
        
        results_data.append({
            "Question": question,
            "Your Answer": user_answer if user_answer else "Not Answered",
            "Correct Answer": correct_answer,
            "Status": "Correct" if is_correct else "Incorrect"
        })
    
    incorrect_count = total_questions - correct_count

    st.subheader("Summary")
    st.markdown(f"**Total Questions:** {total_questions}")
    st.markdown(f"**Correct Answers:** {correct_count}")
    st.markdown(f"**Incorrect Answers:** {incorrect_count}")
    st.markdown(f"**Score:** {correct_count}/{total_questions}")

    st.subheader("Detailed Results")
    st.table(results_data) # Display results in a table

    st.subheader("All Correct Answers (for review)")
    for i, mcq in enumerate(st.session_state.mcqs_for_display):
        question = mcq['question']
        correct_answer = st.session_state.correct_answers_map.get(question)
        st.markdown(f"**Q{i+1}:** {question}")
        st.markdown(f"**Correct Answer:** {correct_answer}")
        st.markdown("---")

elif st.session_state.current_state == "upload":
    st.info("Upload a document and click 'Generate MCQs' to start the test.")

st.markdown("---")
st.markdown("Developed by Hafiz Hassan Abdullah | Email: rajahassan38201@gmail.com | Phone: +92 302 3536363")

